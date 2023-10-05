from abc import ABC, abstractmethod
from typing import List, Tuple, Optional

from PIL import Image

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches

from .constants import IMAGE_FILE_PREFIX, JPG, PDF, DOCX
from .filer_loader import FileLoader
from .quality_reducer import ImageQualityReducer


class ImageLoader(FileLoader):
    async def save_files(self, image_ids: List[str], **kwargs) -> List[str]:
        photo_list = await self._images_processing(image_ids)
        return photo_list

    async def _images_processing(self, image_ids: List[str]) -> List[str]:
        return await self._files_processing(image_ids, IMAGE_FILE_PREFIX, JPG)


class FileConverter(ABC):
    @abstractmethod
    def convert(
        self, conversion_file_path: str, image_paths: List[str], **kwargs
    ) -> None:
        raise NotImplementedError

    @staticmethod
    def compress_check(
        image_file: str, compression_level: Optional[int], **kwargs
    ) -> Tuple[str, int, int]:
        file_extension = kwargs["file_extension"]
        if compression_level:
            img_reducer = ImageQualityReducer(image_file)
            img_obj = img_reducer.reduce(
                quality=compression_level, file_specifier=file_extension
            )
            image_file = img_obj["new_file_path"]
            img = img_obj["reduced_image"]
            width, height = img.size
        else:
            img = Image.open(image_file)
            width, height = img.size

        return image_file, width, height


class PDFConverter(FileConverter):
    def convert(
        self,
        conversion_file_path: str,
        image_paths: List[str],
        compression_level: Optional[int] = None,
    ):
        conversion_canvas = canvas.Canvas(conversion_file_path, pagesize=letter)

        for image_file in image_paths:
            img_data = self.compress(image_file, compression_level)
            image_file, width, height = img_data
            aspect_ratio = height / float(width)
            new_width = letter[0]
            new_height = letter[0] * aspect_ratio
            conversion_canvas.drawImage(
                image_file, 0, 0, width=new_width, height=new_height
            )
            conversion_canvas.showPage()

        conversion_canvas.save()

    def compress(
        self, image_file: str, compression_level: Optional[int]
    ) -> Tuple[str, int, int]:
        return self.compress_check(image_file, compression_level, file_extension=PDF)


class DOCXConverter(FileConverter):
    PHONE_IMAGE_RATIO = 0.5
    DEFAULT_WIDTH_INCHES = Inches(6)
    TALL_IMAGES_HEIGHT_INCHES = Inches(9)

    def convert(
        self,
        conversion_file_path: str,
        image_paths: List[str],
        compression_level: Optional[int] = None,
    ):
        doc = Document()
        for image_file in image_paths:
            img_data = self.compress(image_file, compression_level)
            image_file, width, height = img_data

            if self.is_image_sides_ratio_tall(image_file):
                self.add_picture(doc, image_file, self.TALL_IMAGES_HEIGHT_INCHES)
            else:
                self.add_picture(doc, image_file)

        doc.save(conversion_file_path)

    def add_picture(
        self, doc: Document, image_file: str, height: Optional[Inches] = None
    ) -> None:
        if height:
            doc.add_paragraph().add_run().add_picture(image_file, height=height)
            return None

        doc.add_paragraph().add_run().add_picture(
            image_file, width=self.DEFAULT_WIDTH_INCHES
        )

    def is_image_sides_ratio_tall(self, image_path: str) -> bool:
        with Image.open(image_path) as img:
            aspect_ratio = img.width / img.height
            return aspect_ratio < self.PHONE_IMAGE_RATIO

    def compress(
        self, image_file: str, compression_level: Optional[int]
    ) -> Tuple[str, int, int]:
        return self.compress_check(image_file, compression_level, file_extension=DOCX)
